#!/usr/bin/env python3
from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import re
import sys
from pathlib import Path
from typing import Any


AI_PHRASES = (
    "扎实推进",
    "持续赋能",
    "显著提升",
    "形成闭环",
    "关键路径",
    "抓手",
    "落地见效",
    "稳步推进",
    "阶段性成果",
    "进一步加强",
    "不断优化",
)


def require_docx():
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise SystemExit(
            "Missing dependency: python-docx. Install with `python3 -m pip install python-docx`."
        ) from exc
    return docx


def expand_path(value: str | None) -> Path | None:
    if not value:
        return None
    return Path(os.path.expandvars(os.path.expanduser(value))).resolve()


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx_text(path: Path) -> str:
    docx = require_docx()
    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                ).strip()
                if cell_text:
                    parts.append(cell_text)
    return "\n".join(parts)


def read_source_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return read_docx_text(path)
    return read_text_file(path)


def collect_log_files(log_dir: Path, days: int, require_tag: bool) -> list[Path]:
    if not log_dir.is_dir():
        raise FileNotFoundError(f"log directory does not exist: {log_dir}")

    cutoff = dt.datetime.now() - dt.timedelta(days=days)
    candidates: list[Path] = []
    for path in sorted(log_dir.iterdir()):
        if not path.is_file() or path.suffix.lower() not in {".md", ".txt", ".docx"}:
            continue
        try:
            mtime = dt.datetime.fromtimestamp(path.stat().st_mtime)
        except OSError:
            continue
        if mtime < cutoff:
            continue
        if require_tag:
            try:
                if "#工作日志" not in read_source_text(path):
                    continue
            except Exception:
                continue
        candidates.append(path)
    candidates.sort(key=lambda item: item.stat().st_mtime)
    return candidates


def build_logs_text(log_files: list[Path]) -> str:
    chunks: list[str] = []
    for path in log_files:
        text = read_source_text(path).strip()
        if not text:
            continue
        mtime = dt.datetime.fromtimestamp(path.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
        chunks.append(f"【{mtime}｜{path.name}】\n{text}")
    return "\n\n".join(chunks)


def parse_llm_json(content: str) -> dict[str, Any] | None:
    try:
        data = json.loads(content)
        if isinstance(data, dict):
            return data
    except Exception:
        pass

    start = content.find("{")
    end = content.rfind("}")
    if start == -1 or end <= start:
        return None
    try:
        data = json.loads(content[start : end + 1])
    except Exception:
        return None
    return data if isinstance(data, dict) else None


def get_first_env(names: tuple[str, ...], default: str = "") -> str:
    for name in names:
        value = os.getenv(name)
        if value:
            return value
    return default


def call_llm(logs_text: str) -> dict[str, Any] | None:
    api_key = get_first_env(
        ("WEEKLY_REPORT_API_KEY", "MINIMAX_API_KEY", "OPENAI_API_KEY", "SILICONFLOW_API_KEY")
    )
    if not api_key:
        return None

    try:
        import requests  # type: ignore
    except ImportError:
        return None

    base_url = get_first_env(
        ("WEEKLY_REPORT_BASE_URL", "MINIMAX_BASE_URL", "OPENAI_BASE_URL", "SILICONFLOW_BASE_URL"),
        "https://api.minimaxi.com/v1/chat/completions",
    )
    model = get_first_env(
        ("WEEKLY_REPORT_MODEL", "MINIMAX_MODEL", "OPENAI_MODEL", "SILICONFLOW_MODEL"),
        "MiniMax-M2.5",
    )
    max_tokens = int(get_first_env(("WEEKLY_REPORT_MAX_TOKENS", "MINIMAX_MAX_TOKENS"), "1200"))
    temperature = float(get_first_env(("WEEKLY_REPORT_TEMPERATURE", "MINIMAX_TEMPERATURE"), "0.4"))

    prompt = f"""请根据以下最近工作日志生成中文周报内容。

输出必须是严格 JSON，不要代码块，不要额外解释。结构：
{{"this_week_items":["..."],"summary":"...","next_week_items":["..."],"issues":"..."}}

要求：
1. this_week_items 提炼本周计划/任务，3-10 条，动宾短语。
2. summary 固定以“计划完成情况：”开头，后续按 1、2、3 编号，并与 this_week_items 一一对应。
3. 每条写“对照原计划后的完成情况”，说明做到哪里；没有日志依据时不要编造指标、比例、实验结果或成果。
4. next_week_items 输出 3-8 条下周计划，动宾短语。
5. issues 只写明确问题；没有则写“当前暂无明显问题。”
6. 不要使用这些表达：{", ".join(AI_PHRASES)}。
7. 不要输出 Markdown 标记。

最近工作日志：
{logs_text}
"""
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是一个克制、准确的中文工作周报助手。"},
            {"role": "user", "content": prompt},
        ],
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    try:
        response = requests.post(
            base_url,
            json=payload,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            timeout=120,
        )
        response.raise_for_status()
        content = response.json()["choices"][0]["message"]["content"].strip()
    except Exception as exc:
        print(f"LLM request failed, using fallback summary: {exc}", file=sys.stderr)
        return None

    content = re.sub(r"<think>[\s\S]*?</think>", "", content, flags=re.IGNORECASE).strip()
    return parse_llm_json(content)


def extract_candidate_items(logs_text: str) -> list[str]:
    items: list[str] = []
    for raw_line in logs_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#工作日志"):
            continue
        if line.startswith(("-", "*", "•")):
            item = line.lstrip("-*•").strip()
        elif re.match(r"^\d+[.)、]\s*", line):
            item = re.sub(r"^\d+[.)、]\s*", "", line).strip()
        else:
            continue
        item = re.sub(r"^TODO[:：]\s*", "", item, flags=re.IGNORECASE)
        if item and item not in items:
            items.append(item)
    return items[:10]


def fallback_payload(logs_text: str) -> dict[str, Any]:
    items = extract_candidate_items(logs_text)
    if not items:
        items = ["梳理本周工作记录", "处理当前重点任务", "整理后续推进事项"]

    summary_lines = ["计划完成情况："]
    for idx, item in enumerate(items, start=1):
        summary_lines.append(
            f'{idx}. 对照“{item}”：本周已处理这项工作，具体完成结果需要结合实际材料补充。'
        )

    next_week_items = [
        "梳理本周遗留问题并逐项处理",
        "完善关键任务的验证和记录",
        "整理下周汇报材料和推进计划",
    ]
    return {
        "this_week_items": items,
        "summary": "\n".join(summary_lines),
        "next_week_items": next_week_items,
        "issues": "当前暂无明显问题。",
    }


def normalize_items(value: Any) -> list[str]:
    if isinstance(value, list):
        raw_items = value
    elif isinstance(value, str) and value.strip():
        raw_items = [line.strip() for line in value.splitlines()]
    else:
        raw_items = []
    items: list[str] = []
    for item in raw_items:
        text = str(item).strip()
        text = re.sub(r"^\d+[.)、]\s*", "", text)
        if text:
            items.append(text)
    return items


def format_numbered(items: list[str]) -> str:
    return "\n".join(f"{idx}.{item}" for idx, item in enumerate(items, start=1))


def normalize_docx_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"```[^\n]*\n([\s\S]*?)```", lambda m: m.group(1).strip(), normalized)
    normalized = re.sub(r"`([^`]+)`", r"\1", normalized)
    normalized = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", normalized)
    normalized = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", normalized)
    normalized = re.sub(r"^\s{0,3}#{1,6}\s*", "", normalized, flags=re.MULTILINE)
    normalized = re.sub(r"^\s*>\s?", "", normalized, flags=re.MULTILINE)
    normalized = re.sub(r"(\*\*|__)(.*?)\1", r"\2", normalized)
    normalized = re.sub(r"~~(.*?)~~", r"\1", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def write_cell(cell: Any, text: str) -> None:
    if cell is None:
        return
    if not cell.paragraphs:
        cell.add_paragraph()
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(normalize_docx_text(text))


def find_data_cell(document: Any, header_candidates: tuple[str, ...]) -> Any | None:
    normalized_targets = [target.replace(" ", "") for target in header_candidates]
    for table in document.tables:
        if len(table.rows) < 2:
            continue
        for col_idx, cell in enumerate(table.rows[0].cells):
            cell_text = "".join(p.text.strip() for p in cell.paragraphs).replace(" ", "")
            if any(target in cell_text for target in normalized_targets):
                try:
                    return table.cell(1, col_idx)
                except Exception:
                    return None
    return None


def fill_existing_template(document: Any, payload: dict[str, Any]) -> bool:
    this_week = format_numbered(normalize_items(payload.get("this_week_items")))
    next_week = format_numbered(normalize_items(payload.get("next_week_items")))
    summary = str(payload.get("summary") or "").strip()
    issues = str(payload.get("issues") or "").strip() or extract_issues(summary)

    filled = False
    field_map = [
        (("本周工作计划", "工作计划", "本周计划"), this_week),
        (("主要进展及成果", "计划完成情况", "完成情况", "主要进展"), summary),
        (("存在的问题", "存在问题", "问题"), issues or "当前暂无明显问题。"),
        (("下周工作计划", "下周计划"), next_week),
    ]
    for headers, text in field_map:
        cell = find_data_cell(document, headers)
        if cell is not None:
            write_cell(cell, text)
            filled = True
    return filled


def extract_issues(summary: str) -> str:
    lines = [line.strip() for line in summary.splitlines() if line.strip()]
    issue_lines: list[str] = []
    in_issues = False
    for line in lines:
        if line.startswith(("存在问题：", "存在问题:")):
            in_issues = True
            issue_lines.append(re.split(r"[:：]", line, maxsplit=1)[-1].strip())
            continue
        if in_issues:
            issue_lines.append(line)
    issue_text = "\n".join(line for line in issue_lines if line).strip()
    return issue_text or "当前暂无明显问题。"


def create_default_document(payload: dict[str, Any], title: str) -> Any:
    docx = require_docx()
    document = docx.Document()
    document.add_heading(title, level=1)
    table = document.add_table(rows=2, cols=4)
    headers = ["本周工作计划", "主要进展及成果", "存在的问题", "下周工作计划"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    table.cell(1, 0).text = format_numbered(normalize_items(payload.get("this_week_items")))
    table.cell(1, 1).text = normalize_docx_text(str(payload.get("summary") or ""))
    table.cell(1, 2).text = normalize_docx_text(str(payload.get("issues") or "当前暂无明显问题。"))
    table.cell(1, 3).text = format_numbered(normalize_items(payload.get("next_week_items")))
    return document


def save_report(payload: dict[str, Any], template: Path | None, output_dir: Path, user_name: str, report_date: str) -> Path:
    docx = require_docx()
    output_dir.mkdir(parents=True, exist_ok=True)
    title = f"周工作计划周工作总结-{user_name}{report_date}"
    output_path = output_dir / f"{title}.docx"

    if template:
        document = docx.Document(str(template))
        filled = fill_existing_template(document, payload)
        if not filled:
            document.add_heading("周报内容", level=1)
            document.add_paragraph(normalize_docx_text(str(payload.get("summary") or "")))
    else:
        document = create_default_document(payload, title)

    document.save(str(output_path))
    return output_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a Chinese weekly report DOCX from work logs.")
    parser.add_argument("--log-dir", help="Directory containing recent .md/.txt/.docx work logs.")
    parser.add_argument("--logs-file", help="Single .md/.txt/.docx log file to use instead of --log-dir.")
    parser.add_argument("--days", type=int, default=7, help="Recent day window for --log-dir. Default: 7.")
    parser.add_argument(
        "--include-untagged",
        action="store_true",
        help="Include log files even if they do not contain #工作日志.",
    )
    parser.add_argument("--template", help="Optional DOCX template path. If omitted, a basic DOCX is created.")
    parser.add_argument("--output-dir", default=".", help="Directory for generated DOCX. Default: current directory.")
    parser.add_argument("--user-name", default=os.getenv("USER_NAME", "用户"), help="Name used in the output filename.")
    parser.add_argument(
        "--report-date",
        default=dt.date.today().isoformat(),
        help="Date string used in the output filename. Default: today.",
    )
    parser.add_argument("--no-api", action="store_true", help="Skip LLM API calls and use deterministic fallback text.")
    parser.add_argument("--dry-run", action="store_true", help="Validate inputs and print planned actions without writing DOCX.")
    return parser


def load_logs(args: argparse.Namespace) -> tuple[str, list[Path]]:
    logs_file = expand_path(args.logs_file)
    log_dir = expand_path(args.log_dir)
    if logs_file:
        if not logs_file.is_file():
            raise FileNotFoundError(f"logs file does not exist: {logs_file}")
        return read_source_text(logs_file), [logs_file]
    if log_dir:
        files = collect_log_files(log_dir, args.days, require_tag=not args.include_untagged)
        return build_logs_text(files), files
    return "", []


def main() -> int:
    args = build_parser().parse_args()
    template = expand_path(args.template)
    output_dir = expand_path(args.output_dir) or Path.cwd()

    if template and not template.is_file():
        print(f"template does not exist: {template}", file=sys.stderr)
        return 2

    try:
        logs_text, log_files = load_logs(args)
    except Exception as exc:
        print(f"weekly-report input error: {exc}", file=sys.stderr)
        return 2

    if args.dry_run:
        print(f"log_files={len(log_files)}")
        for path in log_files:
            print(f"- {path}")
        print(f"template={template or '(default document)'}")
        print(f"output_dir={output_dir}")
        print(f"api_mode={'fallback' if args.no_api else 'auto'}")
        return 0

    payload = None if args.no_api else call_llm(logs_text)
    if not payload:
        payload = fallback_payload(logs_text)

    output_path = save_report(
        payload=payload,
        template=template,
        output_dir=output_dir,
        user_name=args.user_name,
        report_date=args.report_date,
    )
    print(f"weekly report generated: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
