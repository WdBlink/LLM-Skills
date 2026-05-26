#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path
from typing import Any

EMU_PER_INCH = 914400
CAPTION_RE = re.compile(r"^\s*([表图])\s*(\d+(?:[-.．]\d+)*)\s+(.+?)\s*$")
PLACEHOLDER_RE = re.compile(r"(XXXX|202X|XX月|插入后删除|本段为|第X|V×|××|XXX|待正式|待补充|待判定|未提供|……)")


def document_text(doc: Any) -> str:
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def table_shapes(doc: Any) -> list[list[int]]:
    return [[len(row.cells) for row in table.rows] for table in doc.tables]


def paragraph_styles(doc: Any) -> list[str]:
    return [p.style.name if p.style is not None else "" for p in doc.paragraphs]


def usable_page_width(doc: Any) -> int:
    widths = [int(section.page_width - section.left_margin - section.right_margin) for section in doc.sections]
    return min(widths) if widths else int(5.8 * EMU_PER_INCH)


def expected_image_count(plan: dict[str, Any]) -> int:
    replacements = plan.get("replacements", [])
    if not isinstance(replacements, list):
        return 0
    return sum(1 for item in replacements if str(item.get("target", "")).startswith("image_"))


def caption_duplicates(doc: Any) -> list[str]:
    seen: dict[tuple[str, str], list[str]] = {}
    for index, paragraph in enumerate(doc.paragraphs):
        style = paragraph.style.name if paragraph.style is not None else ""
        if style.lower().startswith("toc"):
            continue
        match = CAPTION_RE.match(paragraph.text.strip())
        if not match:
            continue
        key = (match.group(1), match.group(2))
        seen.setdefault(key, []).append(f"P{index}:{paragraph.text.strip()}")
    duplicates: list[str] = []
    for key, values in seen.items():
        if len(values) > 1:
            duplicates.append(f"{key[0]} {key[1]} duplicated: {' | '.join(values)}")
    return duplicates


def validate_zip(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
    if bad:
        raise ValueError(f"corrupt DOCX zip member: {bad}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate a filled strict-format DOCX template.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--plan", type=Path)
    parser.add_argument("--strict-layout", action="store_true", help="Fail if paragraph count/styles or section geometry drift.")
    parser.add_argument("--check-placeholders", action="store_true", help="Fail on common unfilled template markers.")
    args = parser.parse_args()

    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit("Missing dependency: install python-docx with `python3 -m pip install python-docx`.") from exc

    validate_zip(args.output)
    template = Document(str(args.template))
    output = Document(str(args.output))

    errors: list[str] = []
    if len(template.tables) != len(output.tables):
        errors.append(f"table count changed: template={len(template.tables)} output={len(output.tables)}")
    if table_shapes(template) != table_shapes(output):
        errors.append("table row/cell shape changed")
    if args.strict_layout:
        if len(template.paragraphs) != len(output.paragraphs):
            errors.append(f"paragraph count changed: template={len(template.paragraphs)} output={len(output.paragraphs)}")
        elif paragraph_styles(template) != paragraph_styles(output):
            errors.append("paragraph style sequence changed")
        for index, (left, right) in enumerate(zip(template.sections, output.sections)):
            left_geometry = (left.page_width, left.page_height, left.left_margin, left.right_margin, left.top_margin, left.bottom_margin)
            right_geometry = (right.page_width, right.page_height, right.left_margin, right.right_margin, right.top_margin, right.bottom_margin)
            if left_geometry != right_geometry:
                errors.append(f"section geometry changed at section {index}")

    required_terms: list[str] = []
    required_images = 0
    if args.plan:
        plan = json.loads(args.plan.read_text(encoding="utf-8"))
        required_terms = [str(term) for term in plan.get("required_terms", [])]
        required_images = expected_image_count(plan)

    text = document_text(output)
    for term in required_terms:
        if term and term not in text:
            errors.append(f"required term missing from output: {term}")
    if required_images and len(output.inline_shapes) < required_images:
        errors.append(f"inline image count is lower than plan image replacements: expected>={required_images} actual={len(output.inline_shapes)}")
    max_width = usable_page_width(output)
    for index, shape in enumerate(output.inline_shapes):
        if shape.width > max_width:
            errors.append(
                f"inline image {index} exceeds usable page width: image={shape.width / EMU_PER_INCH:.2f}in usable={max_width / EMU_PER_INCH:.2f}in"
            )
    duplicates = caption_duplicates(output)
    for duplicate in duplicates:
        errors.append(f"duplicate caption number: {duplicate}")
    if args.check_placeholders:
        markers = sorted(set(PLACEHOLDER_RE.findall(text)))
        if markers:
            errors.append(f"unfilled template markers found: {', '.join(markers)}")

    if errors:
        print("Validation failed:")
        for error in errors:
            print(f"- {error}")
        return 1

    print("Validation passed.")
    print(f"Tables: {len(output.tables)}")
    print(f"Paragraphs: {len(output.paragraphs)}")
    if required_terms:
        print(f"Required terms present: {len(required_terms)}")
    if required_images:
        print(f"Inline images: {len(output.inline_shapes)}")
    print(f"Usable page width: {usable_page_width(output) / EMU_PER_INCH:.2f} in")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
