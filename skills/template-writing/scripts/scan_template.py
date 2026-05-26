#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any


PLACEHOLDER_RE = re.compile(r"(\{\{.+?\}\}|\[.+?\]|_{3,}|[XxＸｘ×]{2,})")
SECTION_HEADING_RE = re.compile(
    r"^("
    r"\d+(\.\d+)*"
    r"|[一二三四五六七八九十]+[、.．]"
    r"|第[一二三四五六七八九十\d]+[章节部分]"
    r")\s*.+"
)
TABLE_OR_FIGURE_CAPTION_RE = re.compile(r"^[表图]\s*\d+([-.．]\d+)*\s+.+")
TOC_LINE_RE = re.compile(r"^.+\t+\d+$|^.+\s{2,}\d+$")
DIAGRAM_KEYWORDS = (
    "结构框图",
    "系统框图",
    "组成框图",
    "结构图",
    "组成图",
    "流程图",
    "架构图",
    "部署图",
    "拓扑图",
    "网络拓扑",
    "关系图",
    "原理图",
    "数据流图",
    "框图",
    "图示",
    "示意图",
)
GUIDANCE_MARKERS = (
    "【要求】",
    "【示例】",
    "填写要求",
    "填写说明",
    "编写要求",
    "编写说明",
    "模板说明",
)


def normalize(text: str) -> str:
    return re.sub(r"[\s\u3000\xa0]+", "", text).strip()


def looks_like_slot(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    return bool(PLACEHOLDER_RE.search(stripped))


def looks_like_guidance(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if any(marker in stripped for marker in GUIDANCE_MARKERS):
        return True
    return stripped in {"要求", "示例", "说明"} or stripped.startswith("注：")


def diagram_keyword(text: str) -> str | None:
    stripped = text.strip()
    if not stripped:
        return None
    for keyword in DIAGRAM_KEYWORDS:
        if keyword in stripped:
            return keyword
    return None


def looks_like_structural_paragraph(text: str, style: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if style.lower().startswith("heading"):
        return True
    if SECTION_HEADING_RE.match(stripped):
        return True
    if TABLE_OR_FIGURE_CAPTION_RE.match(stripped):
        return True
    if TOC_LINE_RE.match(stripped):
        return True
    return False


def paragraph_rewrite_reason(text: str, style: str) -> str | None:
    stripped = text.strip()
    if not stripped:
        return "blank_or_empty_region"
    if diagram_keyword(stripped):
        return "diagram_request"
    if looks_like_guidance(stripped):
        return "template_guidance_or_example"
    if looks_like_slot(stripped):
        return "placeholder_or_slot"
    if TABLE_OR_FIGURE_CAPTION_RE.match(stripped):
        return "caption_text"
    if looks_like_structural_paragraph(stripped, style):
        return None
    return "template_body_text"


def cell_rewrite_reason(row_index: int, text: str) -> str | None:
    stripped = text.strip()
    if not stripped:
        return "blank_or_empty_cell"
    if diagram_keyword(stripped):
        return "diagram_request"
    if looks_like_guidance(stripped):
        return "template_guidance_or_example"
    if looks_like_slot(stripped):
        return "placeholder_or_slot"
    if row_index == 0:
        return None
    return "template_table_content"


def paragraph_record(index: int, paragraph: Any) -> dict[str, Any]:
    text = paragraph.text
    style = paragraph.style.name if paragraph.style is not None else ""
    keyword = diagram_keyword(text)
    return {
        "type": "paragraph",
        "index": index,
        "text": text,
        "style": style,
        "looks_like_slot": looks_like_slot(text),
        "looks_like_guidance": looks_like_guidance(text),
        "looks_like_structure": looks_like_structural_paragraph(text, style),
        "looks_like_diagram_request": keyword is not None,
        "diagram_keyword": keyword,
        "rewrite_reason": paragraph_rewrite_reason(text, style),
    }


def cell_record(table_index: int, row_index: int, col_index: int, cell: Any) -> dict[str, Any]:
    text = cell.text.strip()
    keyword = diagram_keyword(text)
    return {
        "type": "table_cell",
        "table": table_index,
        "row": row_index,
        "col": col_index,
        "text": text,
        "normalized_text": normalize(text),
        "looks_like_slot": looks_like_slot(text),
        "looks_like_guidance": looks_like_guidance(text),
        "looks_like_diagram_request": keyword is not None,
        "diagram_keyword": keyword,
        "rewrite_reason": cell_rewrite_reason(row_index, text),
        "paragraph_count": len(cell.paragraphs),
    }


def adjacent_field_candidates(cells: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    by_row: dict[tuple[int, int], list[dict[str, Any]]] = {}
    for cell in cells:
        by_row.setdefault((cell["table"], cell["row"]), []).append(cell)

    for (_table, _row), row_cells in by_row.items():
        row_cells = sorted(row_cells, key=lambda item: item["col"])
        for left, right in zip(row_cells, row_cells[1:]):
            label = left["normalized_text"]
            if not label or len(label) > 40:
                continue
            if right["looks_like_slot"] or not right["text"]:
                candidates.append(
                    {
                        "label": label,
                        "label_cell": {
                            "table": left["table"],
                            "row": left["row"],
                            "col": left["col"],
                        },
                        "value_cell": {
                            "table": right["table"],
                            "row": right["row"],
                            "col": right["col"],
                        },
                        "current_value": right["text"],
                    }
                )
    return candidates


def extract_guidance_blocks(paragraphs: list[dict[str, Any]], limit: int = 6) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for item in paragraphs:
        if not item["looks_like_guidance"]:
            continue
        start = int(item["index"])
        end = start
        lines = [str(item["text"]).strip()]
        for next_item in paragraphs[start + 1 : start + 1 + limit]:
            text = str(next_item["text"]).strip()
            if not text:
                continue
            if next_item["looks_like_guidance"] and len(lines) > 1:
                break
            lines.append(text)
            end = int(next_item["index"])
        blocks.append(
            {
                "start_paragraph": start,
                "end_paragraph": end,
                "style": item.get("style", ""),
                "marker": str(item["text"]).strip(),
                "text": "\n".join(lines),
            }
        )
    return blocks


def rewrite_candidates(paragraphs: list[dict[str, Any]], cells: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for paragraph in paragraphs:
        reason = paragraph.get("rewrite_reason")
        if reason is None:
            continue
        candidates.append(
            {
                "target": "paragraph",
                "index": paragraph["index"],
                "reason": reason,
                "text": paragraph["text"],
            }
        )
    for cell in cells:
        reason = cell.get("rewrite_reason")
        if reason is None:
            continue
        candidates.append(
            {
                "target": "table_cell",
                "table": cell["table"],
                "row": cell["row"],
                "col": cell["col"],
                "reason": reason,
                "text": cell["text"],
            }
        )
    return candidates


def diagram_candidates(paragraphs: list[dict[str, Any]], cells: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for paragraph in paragraphs:
        if not paragraph.get("looks_like_diagram_request"):
            continue
        candidates.append(
            {
                "target_hint": "image_paragraph or image_before_paragraph/image_after_paragraph",
                "location": {
                    "type": "paragraph",
                    "index": paragraph["index"],
                },
                "diagram_keyword": paragraph.get("diagram_keyword"),
                "text": paragraph["text"],
                "is_caption": bool(TABLE_OR_FIGURE_CAPTION_RE.match(str(paragraph["text"]).strip())),
            }
        )
    for cell in cells:
        if not cell.get("looks_like_diagram_request"):
            continue
        candidates.append(
            {
                "target_hint": "image_table_cell",
                "location": {
                    "type": "table_cell",
                    "table": cell["table"],
                    "row": cell["row"],
                    "col": cell["col"],
                },
                "diagram_keyword": cell.get("diagram_keyword"),
                "text": cell["text"],
                "is_caption": False,
            }
        )
    return candidates


def scan(template: Path) -> dict[str, Any]:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit("Missing dependency: install python-docx with `python3 -m pip install python-docx`.") from exc

    doc = Document(str(template))
    paragraphs = [paragraph_record(i, p) for i, p in enumerate(doc.paragraphs)]

    cells: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    for ti, table in enumerate(doc.tables):
        table_info = {
            "table": ti,
            "rows": len(table.rows),
            "cols_by_row": [len(row.cells) for row in table.rows],
        }
        tables.append(table_info)
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cells.append(cell_record(ti, ri, ci, cell))

    return {
        "template": str(template),
        "template_policy": {
            "mode": "declared_template",
            "rewrite_default": "rewrite all non-structural template content from reference materials",
            "preserve": [
                "section outline and numbering unless the user asks otherwise",
                "table layout and header structure",
                "font, paragraph, page, header, footer, and signature formatting",
            ],
        },
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
        "cells": cells,
        "field_candidates": adjacent_field_candidates(cells),
        "guidance_blocks": extract_guidance_blocks(paragraphs),
        "diagram_candidates": diagram_candidates(paragraphs, cells),
        "rewrite_candidates": rewrite_candidates(paragraphs, cells),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Scan a strict-format DOCX template.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    data = scan(args.template)
    text = json.dumps(data, ensure_ascii=False, indent=2)
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
