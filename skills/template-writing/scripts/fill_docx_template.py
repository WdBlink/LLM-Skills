#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

EMU_PER_INCH = 914400


def clear_paragraph(paragraph: Any) -> None:
    for run in paragraph.runs:
        run.text = ""


def write_text_to_paragraph(paragraph: Any, text: str) -> None:
    runs = paragraph.runs
    if runs:
        first = runs[0]
        clear_paragraph(paragraph)
    else:
        first = paragraph.add_run()
    lines = str(text).splitlines() or [""]
    first.text = lines[0]
    for line in lines[1:]:
        first.add_break()
        first.add_text(line)


def write_text_to_cell(cell: Any, text: str) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    write_text_to_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        clear_paragraph(extra)


def usable_page_width(doc: Any) -> int:
    widths: list[int] = []
    for section in doc.sections:
        widths.append(int(section.page_width - section.left_margin - section.right_margin))
    return min(widths) if widths else int(5.8 * EMU_PER_INCH)


def picture_width(replacement: dict[str, Any], doc: Any | None = None) -> Any:
    allow_overflow = bool(replacement.get("allow_overflow", False))
    usable_width = usable_page_width(doc) if doc is not None else None
    if "width_inches" in replacement:
        from docx.shared import Inches

        width = Inches(float(replacement["width_inches"]))
        if usable_width is not None and not allow_overflow:
            width = min(width, usable_width)
        return width
    if "width_cm" in replacement:
        from docx.shared import Cm

        width = Cm(float(replacement["width_cm"]))
        if usable_width is not None and not allow_overflow:
            width = min(width, usable_width)
        return width
    if usable_width is not None:
        ratio = float(replacement.get("page_width_ratio", 0.96))
        return int(usable_width * max(0.1, min(ratio, 1.0)))
    return None


def write_image_to_paragraph(paragraph: Any, image_path: Path, width: Any = None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    if width is None:
        run.add_picture(str(image_path))
    else:
        run.add_picture(str(image_path), width=width)


def write_image_to_cell(cell: Any, image_path: Path, width: Any = None) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    write_image_to_paragraph(paragraph, image_path, width)
    for extra in cell.paragraphs[1:]:
        clear_paragraph(extra)


def insert_paragraph_near(paragraph: Any, position: str) -> Any:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    if position == "before":
        paragraph._p.addprevious(new_p)
    elif position == "after":
        paragraph._p.addnext(new_p)
    else:
        raise ValueError(f"unsupported insert position: {position}")
    return Paragraph(new_p, paragraph._parent)


def all_paragraphs(doc: Any) -> list[Any]:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def replace_match(doc: Any, match: str, text: str) -> int:
    count = 0
    for paragraph in all_paragraphs(doc):
        for run in paragraph.runs:
            if match in run.text:
                run.text = run.text.replace(match, text)
                count += 1
    return count


def apply_replacement(doc: Any, replacement: dict[str, Any]) -> str:
    target = replacement.get("target")
    selector = replacement.get("selector")
    if selector and not target:
        target = selector.get("type")
        replacement = {**selector, **replacement}

    text = str(replacement.get("text", ""))

    if target == "paragraph":
        index = int(replacement["index"])
        write_text_to_paragraph(doc.paragraphs[index], text)
        return f"paragraph:{index}"

    if target == "paragraph_block":
        start = int(replacement["start"])
        if "end" in replacement:
            end = int(replacement["end"])
        else:
            end = start + int(replacement.get("count", 1)) - 1
        if end < start:
            raise ValueError(f"invalid paragraph_block range: {start}..{end}")
        write_text_to_paragraph(doc.paragraphs[start], text)
        for index in range(start + 1, end + 1):
            clear_paragraph(doc.paragraphs[index])
        return f"paragraph_block:{start}:{end}"

    if target == "table_cell":
        table = int(replacement["table"])
        row = int(replacement["row"])
        col = int(replacement["col"])
        write_text_to_cell(doc.tables[table].rows[row].cells[col], text)
        return f"table_cell:{table}:{row}:{col}"

    if target == "text_match":
        match = str(replacement["match"])
        count = replace_match(doc, match, text)
        if count == 0:
            raise ValueError(f"text_match not found: {match}")
        return f"text_match:{match}:{count}"

    if target in {"image_paragraph", "image_before_paragraph", "image_after_paragraph", "image_table_cell"}:
        image_path = Path(str(replacement["image_path"])).expanduser()
        if not image_path.exists():
            raise FileNotFoundError(f"image_path not found: {image_path}")
        width = picture_width(replacement, doc)

        if target == "image_paragraph":
            index = int(replacement["index"])
            write_image_to_paragraph(doc.paragraphs[index], image_path, width)
            return f"image_paragraph:{index}:{image_path}"

        if target == "image_before_paragraph":
            index = int(replacement["index"])
            paragraph = insert_paragraph_near(doc.paragraphs[index], "before")
            write_image_to_paragraph(paragraph, image_path, width)
            return f"image_before_paragraph:{index}:{image_path}"

        if target == "image_after_paragraph":
            index = int(replacement["index"])
            paragraph = insert_paragraph_near(doc.paragraphs[index], "after")
            write_image_to_paragraph(paragraph, image_path, width)
            return f"image_after_paragraph:{index}:{image_path}"

        table = int(replacement["table"])
        row = int(replacement["row"])
        col = int(replacement["col"])
        write_image_to_cell(doc.tables[table].rows[row].cells[col], image_path, width)
        return f"image_table_cell:{table}:{row}:{col}:{image_path}"

    raise ValueError(f"unsupported replacement target: {target!r}")


def resolve_plan_paths(replacements: list[Any], plan_dir: Path) -> list[Any]:
    resolved: list[Any] = []
    for item in replacements:
        if not isinstance(item, dict) or "image_path" not in item:
            resolved.append(item)
            continue
        image_path = Path(str(item["image_path"])).expanduser()
        if not image_path.is_absolute() and not image_path.exists():
            candidate = plan_dir / image_path
            if candidate.exists():
                item = {**item, "image_path": str(candidate)}
        resolved.append(item)
    return resolved


def main() -> int:
    parser = argparse.ArgumentParser(description="Apply a JSON report plan to a DOCX template.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--plan", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit("Missing dependency: install python-docx with `python3 -m pip install python-docx`.") from exc

    plan = json.loads(args.plan.read_text(encoding="utf-8"))
    replacements = plan.get("replacements", [])
    if not isinstance(replacements, list):
        raise SystemExit("report plan must contain a list named 'replacements'")
    replacements = resolve_plan_paths(replacements, args.plan.parent)

    doc = Document(str(args.template))
    applied = [apply_replacement(doc, item) for item in replacements]

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))

    print(f"Saved {args.output}")
    print(f"Applied {len(applied)} replacements")
    for item in applied:
        print(f"- {item}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
